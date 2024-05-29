from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def set_run_properties(run, properties):
    if 'bold' in properties:
        run.bold = properties['bold']
    if 'italic' in properties:
        run.italic = properties['italic']
    if 'underline' in properties:
        run.underline = properties['underline']
    if 'font_size' in properties:
        run.font.size = Pt(properties['font_size'])
    if 'font_name' in properties:
        run.font.name = properties['font_name']

def set_paragraph_properties(paragraph, properties):
    if 'alignment' in properties:
        alignment = properties['alignment']
        if alignment == 'CENTER (1)':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add more alignment cases as needed
    if 'style' in properties:
        try:
            paragraph.style = properties['style']
        except KeyError:
            paragraph.style = 'Normal'  # Fallback to 'Normal' if the style is not found

def dict_to_docx(doc_dict, docx_path):
    doc = Document()
    
    for para_dict in doc_dict['paragraphs']:
        paragraph = doc.add_paragraph()
        set_paragraph_properties(paragraph, para_dict['properties'])
        
        for run_dict in para_dict['runs']:
            run = paragraph.add_run(run_dict['text'])
            set_run_properties(run, run_dict['properties'])
    
    doc.save(docx_path)

# Dictionary to be converted to DOCX
document_dict = {
    'paragraphs': [
        {'properties': {'style': 'Title'},
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0,
                                  'underline': True},
                   'text': 'Mục lục:'}],
         'text': 'Mục lục:'},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ''}],
         'text': ''},
        {'properties': {'alignment': 'CENTER (1)', 'style': 'Title'},
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 24.0},
                   'text': 'Hướng dẫn chạy script tổng hợp thông tin các phần mềm cài đặt trên máy '}],
         'text': 'Hướng dẫn chạy script tổng hợp thông tin các phần mềm cài đặt trên máy '},
        {'properties': {'style': 'Heading 1'},
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman'},
                   'text': 'Đối với Window:'}],
         'text': 'Đối với Window:'},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'Bước 1: '},
                  {'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'Truy cập vào folder dưới đây:'}],
         'text': 'Bước 1: Truy cập vào folder dưới đây:'},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {}, 'text': 'old_text'}],
         'text': 'old_text'},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [],
         'text': 'Script'},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'Trong thư mục mới hiện ra 3 script, bạn download file script'},
                  {'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ' '},
                  {'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'window-docker.ISO.HN.ps1'},
                  {'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ' về và copy sang thư mục '},
                  {'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'C:\\Users\\Public'},
                  {'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ' trong ổ C. '}],
         'text': 'Trong thư mục mới hiện ra 3 script, bạn download file script window-docker.ISO.HN.ps1 về và copy sang thư mục C:\\Users\\Public trong ổ C. '},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0,
                                  'underline': True},
                   'text': '*Lưu ý:'},
                  {'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ' '},
                  {'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'nếu như các bạn đã có sẵn file cũ trong thư mục này. Hãy xóa file cũ đi và copy file mới vào. Tránh để file có đuôi có số (1), (2), ...'}],
         'text': '*Lưu ý: nếu như các bạn đã có sẵn file cũ trong thư mục này. Hãy xóa file cũ đi và copy file mới vào. Tránh để file có đuôi có số (1), (2), ...'},
        {'properties': {'alignment': 'CENTER (1)',
                        'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {}, 'text': ''}],
         'text': ''},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ''}],
         'text': ''},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ''}],
         'text': ''},
        {'properties': {'style': 'Normal'},  # Use 'Normal' instead of 'LO-normal'
         'runs': [{'properties': {'bold': True,
                                  'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': 'Bước 2:'},
                  {'properties': {'font_name': 'Times New Roman',
                                  'font_size': 12.0},
                   'text': ' Chạy script:'}],
         'text': 'Bước 2: Chạy script:'}
    ]
}

# Path to the new DOCX file
docx_path = 'output_document.docx'
dict_to_docx(document_dict, docx_path)
