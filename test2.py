from docx import Document
from docx.shared import Pt, RGBColor
import json

class DocxTranslator:
    def __init__(self, input_path):
        self.doc = Document(input_path)

    def get_run_attributes(self, run):
        details = {
            "Text": run.text,
            "Bold": run.bold,
            "Italic": run.italic,
            "Underline": run.underline,
            "Font name": run.font.name if run.font else None,
            "Font size": run.font.size.pt if run.font and run.font.size else None,
            "Font color": run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None,
            "Highlight color": run.font.highlight_color if run.font else None,
            "Strike": run.font.strike if run.font else None,
            "Double Strike": run.font.double_strike if run.font else None,
            "Subscript": run.font.subscript if run.font else None,
            "Superscript": run.font.superscript if run.font else None,
            "Style": run.style.name if run.style else 'No style'
        }
        return details

    def append_attributes_paragraph(self, paragraph, run_attributes):
        new_run = paragraph.add_run("\n\n--- Attributes ---\n" + json.dumps(run_attributes, indent=4))
        new_run.font.size = Pt(10)
        new_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color to differentiate attributes text

    def translate_docx(self, output_path):
        # Create a new document for the output
        new_doc = Document()

        # Iterate through each paragraph in the document
        for paragraph in self.doc.paragraphs:
            new_paragraph = new_doc.add_paragraph()
            
            run_attributes_list = []
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font:
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
                    new_run.font.highlight_color = run.font.highlight_color
                    new_run.font.strike = run.font.strike
                    new_run.font.double_strike = run.font.double_strike
                    new_run.font.subscript = run.font.subscript
                    new_run.font.superscript = run.font.superscript
                new_run.style = run.style
                
                if run.text.strip():  # Add only non-empty runs
                    run_attributes_list.append(self.get_run_attributes(run))
                    
            # if run_attributes_list:
            #     self.append_attributes_paragraph(new_paragraph, run_attributes_list)
        
        # Save the new document
        new_doc.save(output_path)
    
    def get_all_attributes(self):
        all_attributes = []
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():  # Add only non-empty runs
                    all_attributes.append(self.get_run_attributes(run))
        return json.dumps(all_attributes, indent=4)

# Example usage
input_path = "TestData/DGP.docx"
output_path = "output_translated_1.docx"

translator = DocxTranslator(input_path)
translator.translate_docx(output_path)

# Get all attributes as a JSON string
attributes_json = translator.get_all_attributes()
print(attributes_json)
