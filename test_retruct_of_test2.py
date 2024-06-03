from docx import Document
from docx.shared import Pt, RGBColor

def get_run_attributes(run):
    details = []
    details.append(f"Text: {run.text}")
    details.append(f"Bold: {run.bold}")
    details.append(f"Italic: {run.italic}")
    details.append(f"Underline: {run.underline}")
    
    if run.font:
        details.append(f"Font name: {run.font.name}")
        details.append(f"Font size: {run.font.size}")
        
        if run.font.color and run.font.color.rgb:
            details.append(f"Font color: {run.font.color.rgb}")
        
        details.append(f"Highlight color: {run.font.highlight_color}")
        details.append(f"Strike: {run.font.strike}")
        details.append(f"Double Strike: {run.font.double_strike}")
        details.append(f"Subscript: {run.font.subscript}")
        details.append(f"Superscript: {run.font.superscript}")
    
    details.append(f"Style: {run.style.name if run.style else 'No style'}")
    
    return "\n".join(details)

def append_attributes_paragraph(paragraph, run_attributes):
    new_run = paragraph.add_run("\n\n--- Attributes ---\n" + run_attributes)
    new_run.font.size = Pt(10)
    new_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color to differentiate attributes text

def copy_paragraph(paragraph, new_doc):
    new_paragraph = new_doc.add_paragraph()
    
    for run in paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if run.font:
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            new_run.font.highlight_color = run.font.highlight_color
            new_run.font.strike = run.font.strike
            new_run.font.double_strike = run.font.double_strike
            new_run.font.subscript = run.font.subscript
            new_run.font.superscript = run.font.superscript
        new_run.style = run.style
    
    run_attributes = "\n\n".join(get_run_attributes(run) for run in paragraph.runs if run.text.strip())
    # append_attributes_paragraph(new_paragraph, run_attributes)

def copy_table(table, new_doc):
    new_table = new_doc.add_table(rows=0, cols=len(table.columns))
    for row in table.rows:
        new_row = new_table.add_row()
        for idx, cell in enumerate(row.cells):
            new_cell = new_row.cells[idx]
            for paragraph in cell.paragraphs:
                copy_paragraph(paragraph, new_cell)

def copy_image(image, new_doc):
    # Add a new paragraph for the image
    new_paragraph = new_doc.add_paragraph()
    run = new_paragraph.add_run()
    run.add_picture(image)

def translate_docx(input_path, output_path):
    # Load the input document
    doc = Document(input_path)
    
    # Create a new document for the output
    new_doc = Document()

    # Iterate through each element in the document
    for element in doc.element.body:
        if element.tag.endswith('p'):
            paragraph = Document().paragraphs[0]._element.__class__(element, doc)
            copy_paragraph(paragraph, new_doc)
        elif element.tag.endswith('tbl'):
            table = Document().tables[0]._element.__class__(element, doc)
            copy_table(table, new_doc)
        elif element.tag.endswith('drawing'):
            copy_image(element, new_doc)

    # Save the new document
    new_doc.save(output_path)

input_path = "TestData/DGP.docx"
output_path = "output_translated.docx"

# Translate the docx file
translate_docx(input_path, output_path)
