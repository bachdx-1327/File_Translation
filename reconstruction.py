from docx import Document
from docx.shared import Inches
import os

# Define paths
image_dir = "img"
output_file = "reconstructed_file.docx"

# Load the content list
content = []
with open(os.path.join(image_dir, 'content.txt'), 'r') as f:
    for line in f:
        item_type, item_value = line.strip().split(':', 1)
        content.append((item_type, item_value))

# Create a new Document
doc = Document()

# Add content to the new Document
for item_type, item_value in content:
    if item_type == 'text':
        doc.add_paragraph(item_value)
    elif item_type == 'image':
        doc.add_picture(item_value, width=Inches(4.0))  # Adjust the width as needed

# Save the new Document
doc.save(output_file)
