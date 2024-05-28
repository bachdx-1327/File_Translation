from docx import Document

def replace_text_with_format(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            new_paragraph_text = ""
            new_runs = []
            for run in paragraph.runs:
                if old_text in run.text:
                    # Split the run text on old_text and insert new_text while preserving formatting
                    parts = run.text.split(old_text)
                    for i, part in enumerate(parts):
                        if i > 0:
                            new_run = paragraph.add_run(new_text)
                            # Copy formatting from original run to new run
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.color:
                                new_run.font.color.rgb = run.font.color.rgb
                            new_run.font.size = run.font.size
                            new_run.font.name = run.font.name
                            new_run.font.highlight_color = run.font.highlight_color
                            new_run.font.strike = run.font.strike
                            new_run.font.superscript = run.font.superscript
                            new_run.font.subscript = run.font.subscript
                            new_run.font.all_caps = run.font.all_caps
                            new_run.font.small_caps = run.font.small_caps
                            new_run.font.shadow = run.font.shadow
                            new_run.font.outline = run.font.outline
                            new_run.font.cs_bold = run.font.cs_bold
                            new_run.font.cs_italic = run.font.cs_italic
                            new_run.font.locale_id = run.font.locale_id
                            new_run.font.language_id = run.font.language_id
                            new_runs.append(new_run)
                        new_runs.append(paragraph.add_run(part))
                else:
                    new_runs.append(run)
            # Clear existing runs and add the modified runs
            paragraph.clear()
            for run in new_runs:
                paragraph.add_run(run.text).font = run.font

# Tạo đối tượng Document từ file Word
doc = Document("TestData/DGP_1.docx")

# Gọi hàm để thay thế văn bản
replace_text_with_format(doc, "old_text", "new_text")

# Lưu lại file Word mới
doc.save("modified_example.docx")
