from docx import Document
from deeptrans import ja2vi

def translate_run_text(run_text):
    return ja2vi(run_text)

def translate_paragraph(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            print(run.text)
            if len(run.text) <= 2: 
                pass
            else:
                translated_text = translate_run_text(run.text)
                run.text = translated_text
            # print(translated_text)
    

def translate_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        print(len(run.text))
                        translated_text = translate_run_text(run.text)
                        run.text = translated_text

def translate_docx(input_path, out_path):
    # Loading the docx file 
    doc = Document(input_path)
    
    # Iterate through each paragraph in the document
    for paragraph in doc.paragraphs:
        translate_paragraph(paragraph)

    # Iterate through each table in the document
    for table in doc.tables:
        translate_table(table)
    
    # Save the translated document
    doc.save(out_path)

input_path = "document_5.docx"
out_path = "document_6.docx"

# Translate the docx file
translate_docx(input_path, out_path)